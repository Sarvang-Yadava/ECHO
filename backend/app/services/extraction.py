from __future__ import annotations

import re
from pathlib import Path

import fitz
from docx import Document as DocxDocument
from PIL import Image

from app.services.ocr import read_image


def extract_document(file_path: Path) -> tuple[str, int, str]:
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return _pdf(file_path)
    if suffix == ".docx":
        return _docx(file_path)
    if suffix in {".png", ".jpg", ".jpeg"}:
        return _image(file_path)
    if suffix == ".txt":
        return _text(file_path)
    raise RuntimeError(f"Unsupported file type: {suffix or 'unknown'}")


def _pdf(file_path: Path) -> tuple[str, int, str]:
    document = fitz.open(file_path)
    pages = []
    poor_pages = 0
    for number, page in enumerate(document, 1):
        content = page.get_text("text").strip()
        # The page marker makes source references and heading recovery durable.
        pages.append(f"[Page {number}]\n{content}")
        if len(re.sub(r"\s+", "", content)) < 40:
            poor_pages += 1
    text = "\n\n".join(pages).strip()
    if poor_pages == 0 and len(re.sub(r"\s+", "", text)) >= 80:
        return text, document.page_count, "pymupdf"

    ocr_pages = []
    for number, page in enumerate(document, 1):
        pix = page.get_pixmap(dpi=220, alpha=False)
        image = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
        ocr_pages.append(f"[Page {number}]\n{read_image(image)}")
    return "\n\n".join(ocr_pages), document.page_count, "ocr-fallback"


def _docx(file_path: Path) -> tuple[str, int, str]:
    document = DocxDocument(file_path)
    blocks: list[str] = []
    for paragraph in document.paragraphs:
        value = paragraph.text.strip()
        if not value:
            continue
        style = (paragraph.style.name or "").lower()
        blocks.append(f"## {value}" if "heading" in style or "title" in style else value)
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                blocks.append(row_text)
    text = "\n".join(blocks).strip()
    if not text:
        raise RuntimeError("The DOCX contains no readable paragraphs or tables")
    return text, max(1, len(blocks) // 18), "docx-structured"


def _image(file_path: Path) -> tuple[str, int, str]:
    with Image.open(file_path) as image:
        return "[Page 1]\n" + read_image(image.convert("RGB")), 1, "ocr-image"


def _text(file_path: Path) -> tuple[str, int, str]:
    text = file_path.read_text(encoding="utf-8", errors="ignore").strip()
    if not text:
        raise RuntimeError("The text file is empty")
    return text, max(1, text.count("\f") + 1), "text"
